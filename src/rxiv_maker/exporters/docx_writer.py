"""DOCX writer for rxiv-maker export.

This module handles the actual generation of DOCX files using python-docx,
writing structured content with formatting, citations, and references.
"""

from pathlib import Path
from typing import Any, Dict, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from lxml import etree

from ..core.logging_config import get_logger
from ..utils.docx_helpers import convert_pdf_to_image

logger = get_logger()


class DocxWriter:
    """Writes structured content to DOCX files using python-docx."""

    def write(
        self,
        doc_structure: Dict[str, Any],
        bibliography: Dict[int, Dict],
        output_path: Path,
        include_footnotes: bool = True,
        base_path: Optional[Path] = None,
    ) -> Path:
        """Write DOCX file from structured content.

        Args:
            doc_structure: Structured document with sections
            bibliography: Bibliography entries mapped by number
            output_path: Path where DOCX file should be saved
            include_footnotes: Whether to add DOI footnotes
            base_path: Base path for resolving relative figure paths

        Returns:
            Path to created DOCX file
        """
        self.base_path = base_path or Path.cwd()
        self.bibliography = bibliography
        self.include_footnotes = include_footnotes
        doc = Document()

        # Collect figures to add at the end and create label->number mapping
        figures = []
        figure_map = {}  # Maps label to number

        for section in doc_structure["sections"]:
            if section["type"] == "figure":
                label = section.get("label", "")
                if label:
                    figure_map[label] = len(figures) + 1  # 1-indexed
                figures.append(section)

        # Store figure map for use in text processing
        self.figure_map = figure_map

        # Process each section (skip figures for now)
        for section in doc_structure["sections"]:
            if section["type"] == "figure":
                continue  # Skip figures, we'll add them at the end
            else:
                self._add_section(doc, section, bibliography, include_footnotes)

        # Add bibliography section (before figures)
        if include_footnotes and bibliography:
            doc.add_page_break()
            doc.add_heading("Bibliography", level=1)

            # Add numbered bibliography entries
            for num in sorted(bibliography.keys()):
                bib_entry = bibliography[num]
                para = doc.add_paragraph()

                # Add citation number in bold
                num_run = para.add_run(f"[{num}] ")
                num_run.bold = True

                # Add formatted bibliography text (slim format)
                para.add_run(bib_entry["formatted"])

                # Add DOI as hyperlink if present
                if bib_entry.get("doi"):
                    doi = bib_entry["doi"]
                    doi_url = f"https://doi.org/{doi}" if not doi.startswith("http") else doi
                    para.add_run(" ")
                    self._add_hyperlink(para, doi_url, doi_url)

                # Add spacing between entries
                para.paragraph_format.space_after = Pt(6)

        # Add figures at the end
        if figures:
            # Add page break before figures
            doc.add_page_break()

            # Add "Figures" heading
            doc.add_heading("Figures", level=1)

            # Add each figure with its caption and number
            for i, figure in enumerate(figures, start=1):
                self._add_figure(doc, figure, figure_number=i)
                # Spacing is handled within _add_figure (space_after on caption)

        # Save document
        doc.save(str(output_path))
        return output_path

    def _add_section(
        self,
        doc: Document,
        section: Dict[str, Any],
        bibliography: Dict[int, Dict],
        include_footnotes: bool,
    ):
        """Add a section to the document.

        Args:
            doc: Document object
            section: Section data
            bibliography: Bibliography entries
            include_footnotes: Whether to add footnotes
        """
        section_type = section["type"]

        if section_type == "heading":
            self._add_heading(doc, section)
        elif section_type == "paragraph":
            self._add_paragraph(doc, section, bibliography, include_footnotes)
        elif section_type == "list":
            self._add_list(doc, section)
        elif section_type == "code_block":
            self._add_code_block(doc, section)
        elif section_type == "figure":
            self._add_figure(doc, section)

    def _add_heading(self, doc: Document, section: Dict[str, Any]):
        """Add heading to document.

        Args:
            doc: Document object
            section: Heading section data with 'level' and 'text'
        """
        level = section["level"]
        text = section["text"]
        doc.add_heading(text, level=level)

    def _add_paragraph(
        self,
        doc: Document,
        section: Dict[str, Any],
        bibliography: Dict[int, Dict],
        include_footnotes: bool,
    ):
        """Add paragraph with formatted runs to document.

        Args:
            doc: Document object
            section: Paragraph section data with 'runs'
            bibliography: Bibliography entries
            include_footnotes: Whether to add footnotes
        """
        paragraph = doc.add_paragraph()

        # Set justified alignment for all paragraphs
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        runs_data = section["runs"]

        for run_data in runs_data:
            self._add_run(paragraph, run_data, bibliography, include_footnotes)

    def _add_run(self, paragraph, run_data: Dict[str, Any], bibliography: Dict[int, Dict], include_footnotes: bool):
        """Add a single run to a paragraph.

        Args:
            paragraph: Paragraph object
            run_data: Run data with type and formatting
            bibliography: Bibliography entries
            include_footnotes: Whether to add footnotes
        """
        if run_data["type"] == "text":
            text = run_data["text"]
            run = paragraph.add_run(text)

            # Apply formatting
            if run_data.get("bold"):
                run.bold = True
            if run_data.get("italic"):
                run.italic = True
            if run_data.get("code"):
                run.font.name = "Courier New"
                run.font.size = Pt(10)
            if run_data.get("xref"):
                # Cross-reference - apply yellow highlighting
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW

        elif run_data["type"] == "citation":
            cite_num = run_data["number"]
            # Add citation as [NN] inline with yellow highlighting
            run = paragraph.add_run(f"[{cite_num}]")
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            run.font.size = Pt(10)

    def _add_list(self, doc: Document, section: Dict[str, Any]):
        """Add list to document.

        Args:
            doc: Document object
            section: List section data with 'list_type' and 'items'
        """
        list_type = section["list_type"]
        items = section["items"]

        for item in items:
            doc.add_paragraph(item, style="List Bullet" if list_type == "bullet" else "List Number")

    def _add_code_block(self, doc: Document, section: Dict[str, Any]):
        """Add code block to document.

        Args:
            doc: Document object
            section: Code block section data with 'content'
        """
        code_content = section["content"]
        paragraph = doc.add_paragraph(code_content)

        # Style as code
        for run in paragraph.runs:
            run.font.name = "Courier New"
            run.font.size = Pt(9)

        # Set paragraph formatting
        paragraph_format = paragraph.paragraph_format
        paragraph_format.left_indent = Pt(36)  # Indent code blocks

    def _add_figure(self, doc: Document, section: Dict[str, Any], figure_number: int = None):
        """Add figure to document with caption.

        Args:
            doc: Document object
            section: Figure section data with 'path', 'caption_runs', 'label'
            figure_number: Figure number (1-indexed)
        """
        figure_path = Path(section["path"])
        caption_runs = section.get("caption_runs", [])

        # Resolve relative path
        if not figure_path.is_absolute():
            figure_path = self.base_path / figure_path

        # Try to convert PDF to image if it's a PDF
        img_bytes = None
        if figure_path.exists() and figure_path.suffix.lower() == ".pdf":
            img_bytes = convert_pdf_to_image(figure_path)
        elif not figure_path.exists():
            logger.warning(f"Figure file not found: {figure_path}")

        if img_bytes:
            # Add image
            try:
                doc.add_picture(img_bytes, width=Inches(6))
                logger.debug(f"Embedded figure: {figure_path}")
            except Exception as e:
                logger.warning(f"Failed to embed figure {figure_path}: {e}")
                # Add placeholder text
                p = doc.add_paragraph()
                run = p.add_run(f"[Figure: {figure_path.name}]")
                run.italic = True
        else:
            # Add placeholder if conversion failed or not a PDF
            p = doc.add_paragraph()
            run = p.add_run(f"[Figure: {figure_path.name}]")
            run.italic = True
            logger.warning(f"Could not embed figure: {figure_path}")

        # Add caption with formatted runs
        if caption_runs:
            caption_para = doc.add_paragraph()
            caption_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            # Format as "Figure number: caption"
            if figure_number:
                run = caption_para.add_run(f"Figure {figure_number}: ")
                run.bold = True
            else:
                run = caption_para.add_run("Figure: ")
                run.bold = True

            # Add formatted caption runs
            for run_data in caption_runs:
                self._add_run(caption_para, run_data, {}, False)

            # Add spacing after figure
            caption_para.paragraph_format.space_after = Pt(12)

    def _add_footnote(self, paragraph, cite_num: int, bib_entry: Dict[str, Any]):
        """Add footnote with bibliography entry and DOI.

        Args:
            paragraph: Paragraph object to attach footnote to
            cite_num: Citation number
            bib_entry: Bibliography entry with 'formatted' text and optional 'doi'
        """
        # Ensure footnotes part exists
        self._ensure_footnotes_part(paragraph.part)

        # Create footnote reference in text
        run = paragraph.add_run()

        # Add footnote reference element
        footnote_ref = OxmlElement("w:footnoteReference")
        footnote_ref.set(qn("w:id"), str(cite_num))
        run._element.append(footnote_ref)

        # Get footnotes part
        footnotes_part = paragraph.part.part_related_by(
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes"
        )

        # Create footnote element with hyperlink
        formatted_text = bib_entry["formatted"]
        doi = bib_entry.get("doi")

        footnote_elem = self._create_footnote_element(cite_num, formatted_text, doi, footnotes_part)

        # Append to footnotes
        footnotes_part.element.append(footnote_elem)

    def _ensure_footnotes_part(self, doc_part):
        """Ensure the document has a footnotes part with required separators."""
        try:
            # Check if it exists
            doc_part.part_related_by("http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes")
        except KeyError:
            # Create footnotes part
            from docx.opc.constants import RELATIONSHIP_TYPE as RT
            from docx.opc.packuri import PackURI
            from docx.opc.part import XmlPart

            # Create proper XML with lxml
            footnotes_xml = etree.Element(
                qn("w:footnotes"),
                nsmap={
                    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
                    "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
                },
            )

            # Add separator footnote (required by Word)
            separator = etree.SubElement(footnotes_xml, qn("w:footnote"))
            separator.set(qn("w:type"), "separator")
            separator.set(qn("w:id"), "-1")
            sep_p = etree.SubElement(separator, qn("w:p"))
            sep_r = etree.SubElement(sep_p, qn("w:r"))
            _sep_t = etree.SubElement(sep_r, qn("w:separator"))  # noqa: F841

            # Add continuation separator footnote (required by Word)
            cont_sep = etree.SubElement(footnotes_xml, qn("w:footnote"))
            cont_sep.set(qn("w:type"), "continuationSeparator")
            cont_sep.set(qn("w:id"), "0")
            cont_p = etree.SubElement(cont_sep, qn("w:p"))
            cont_r = etree.SubElement(cont_p, qn("w:r"))
            _cont_t = etree.SubElement(cont_r, qn("w:continuationSeparator"))  # noqa: F841

            # Create the part
            partname = PackURI("/word/footnotes.xml")
            content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml"
            footnotes_part = XmlPart(partname, content_type, footnotes_xml, doc_part.package)

            # Add relationship
            doc_part.relate_to(footnotes_part, RT.FOOTNOTES)

    def _create_footnote_element(self, footnote_id: int, text: str, doi: str | None, footnotes_part) -> etree.Element:
        """Create a footnote XML element.

        Args:
            footnote_id: Footnote ID number
            text: Bibliography text
            doi: Optional DOI URL
            footnotes_part: Footnotes part for creating hyperlink relationships

        Returns:
            lxml Element for footnote
        """
        # Create footnote element
        footnote = etree.Element(qn("w:footnote"))
        footnote.set(qn("w:id"), str(footnote_id))

        # Create paragraph in footnote
        p = etree.SubElement(footnote, qn("w:p"))

        # Add paragraph properties with footnote style
        pPr = etree.SubElement(p, qn("w:pPr"))
        pStyle = etree.SubElement(pPr, qn("w:pStyle"))
        pStyle.set(qn("w:val"), "FootnoteText")

        # Add footnote reference mark
        r = etree.SubElement(p, qn("w:r"))
        rPr = etree.SubElement(r, qn("w:rPr"))
        rStyle = etree.SubElement(rPr, qn("w:rStyle"))
        rStyle.set(qn("w:val"), "FootnoteReference")
        etree.SubElement(r, qn("w:footnoteRef"))

        # Add bibliography text (excluding DOI since we'll add it separately as hyperlink)
        # Remove DOI from text if present
        text_without_doi = text.split("\nDOI:")[0] if "\nDOI:" in text else text

        r2 = etree.SubElement(p, qn("w:r"))
        # Add font size 8pt
        rPr2 = etree.SubElement(r2, qn("w:rPr"))
        sz = etree.SubElement(rPr2, qn("w:sz"))
        sz.set(qn("w:val"), "16")  # 16 half-points = 8pt
        szCs = etree.SubElement(rPr2, qn("w:szCs"))
        szCs.set(qn("w:val"), "16")  # For complex scripts

        t = etree.SubElement(r2, qn("w:t"))
        t.set(qn("xml:space"), "preserve")
        t.text = f" {text_without_doi}"

        # Add DOI as clickable hyperlink if present
        if doi:
            doi_url = f"https://doi.org/{doi}" if not doi.startswith("http") else doi

            # Create relationship for hyperlink in footnotes part
            r_id = footnotes_part.relate_to(
                doi_url,
                "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
                is_external=True,
            )

            # Create hyperlink element
            hyperlink = etree.SubElement(p, qn("w:hyperlink"))
            hyperlink.set(qn("r:id"), r_id)

            # Create run with link text
            r3 = etree.SubElement(hyperlink, qn("w:r"))

            # Add run properties for hyperlink style
            rPr3 = etree.SubElement(r3, qn("w:rPr"))
            rStyle3 = etree.SubElement(rPr3, qn("w:rStyle"))
            rStyle3.set(qn("w:val"), "Hyperlink")

            # Add font size 8pt
            sz3 = etree.SubElement(rPr3, qn("w:sz"))
            sz3.set(qn("w:val"), "16")  # 16 half-points = 8pt
            szCs3 = etree.SubElement(rPr3, qn("w:szCs"))
            szCs3.set(qn("w:val"), "16")

            # Add underline
            u = etree.SubElement(rPr3, qn("w:u"))
            u.set(qn("w:val"), "single")

            # Add color (blue)
            color = etree.SubElement(rPr3, qn("w:color"))
            color.set(qn("w:val"), "0000FF")

            # Add text (just the URL, no "DOI:" prefix)
            t2 = etree.SubElement(r3, qn("w:t"))
            t2.set(qn("xml:space"), "preserve")
            t2.text = f" {doi_url}"

        return footnote

    def _add_hyperlink(self, paragraph, url: str, text: str):
        """Add hyperlink to paragraph.

        Args:
            paragraph: Paragraph object
            url: URL to link to
            text: Display text
        """
        # Add hyperlink using OOXML
        part = paragraph.part
        r_id = part.relate_to(
            url,
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
            is_external=True,
        )

        # Create hyperlink element
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)

        # Create run element
        new_run = OxmlElement("w:r")

        # Create run properties
        r_pr = OxmlElement("w:rPr")

        # Add underline
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        r_pr.append(u)

        # Add color (blue)
        color = OxmlElement("w:color")
        color.set(qn("w:val"), "0000FF")
        r_pr.append(color)

        new_run.append(r_pr)

        # Create text element
        text_element = OxmlElement("w:t")
        text_element.text = text
        new_run.append(text_element)

        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)
