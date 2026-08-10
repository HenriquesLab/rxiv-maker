"""Mark revision changes between two DOCX exports.

Journals that ask for a "highlighted" revision want a Word file showing what
moved between submission and revision, and several explicitly reject Word's
Track Changes because the markup is lost when they convert to PDF. This module
produces that artefact directly: insertions highlighted, deletions struck
through and coloured.

The comparison runs on the two rendered DOCX files rather than on the Markdown
sources, so what is marked is what a reader actually sees.
"""

import difflib
import re
from pathlib import Path
from typing import List, Optional, Tuple

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import RGBColor

from ..core.logging_config import get_logger

logger = get_logger()

# Deletions are struck through in red; insertions keep the document font and are
# highlighted instead, which survives PDF conversion.
DELETED_COLOR = RGBColor(0xC0, 0x00, 0x00)
INSERTED_HIGHLIGHT = WD_COLOR_INDEX.BRIGHT_GREEN
DELETED_HIGHLIGHT = WD_COLOR_INDEX.PINK

# Split into words while keeping the whitespace, so rebuilt text keeps its spacing.
_TOKEN = re.compile(r"\S+\s*")

# Below this similarity, two paragraphs paired by the diff are treated as
# unrelated: interleaving their words would produce an unreadable hybrid, so the
# old one is shown deleted and the new one inserted.
SIMILARITY_THRESHOLD = 0.5


def _tokenize(text: str) -> List[str]:
    """Split text into whitespace-preserving word tokens."""
    return _TOKEN.findall(text)


def _similarity(old_text: str, new_text: str) -> float:
    """Return how alike two paragraphs are, between 0 and 1."""
    return difflib.SequenceMatcher(None, old_text, new_text, autojunk=False).ratio()


def _iter_block_paragraphs(doc) -> List:
    """Return the document's paragraphs, including those inside tables."""
    paragraphs = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.extend(cell.paragraphs)
    return paragraphs


def _clear_runs(paragraph) -> None:
    """Remove every run from a paragraph, leaving it empty."""
    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)


def _add_marked_run(paragraph, text: str, mode: str) -> None:
    """Append a run marked as unchanged, inserted, or deleted."""
    if not text:
        return
    run = paragraph.add_run(text)
    if mode == "insert":
        run.font.highlight_color = INSERTED_HIGHLIGHT
    elif mode == "delete":
        run.font.strike = True
        run.font.color.rgb = DELETED_COLOR
        run.font.highlight_color = DELETED_HIGHLIGHT


def _mark_paragraph(paragraph, old_text: str, new_text: str) -> None:
    """Rewrite a paragraph as a word-level diff of old vs new text.

    Existing runs are replaced, so inline formatting inside a changed paragraph
    is not preserved. The clean export carries the real formatting; this file
    exists to show what changed.
    """
    old_tokens, new_tokens = _tokenize(old_text), _tokenize(new_text)
    matcher = difflib.SequenceMatcher(None, old_tokens, new_tokens, autojunk=False)

    _clear_runs(paragraph)
    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == "equal":
            _add_marked_run(paragraph, "".join(new_tokens[j1:j2]), "equal")
        elif tag == "insert":
            _add_marked_run(paragraph, "".join(new_tokens[j1:j2]), "insert")
        elif tag == "delete":
            _add_marked_run(paragraph, "".join(old_tokens[i1:i2]), "delete")
        elif tag == "replace":
            _add_marked_run(paragraph, "".join(old_tokens[i1:i2]), "delete")
            _add_marked_run(paragraph, "".join(new_tokens[j1:j2]), "insert")


def _insert_deleted_paragraph(anchor, text: str) -> None:
    """Insert a fully struck-through paragraph before the anchor paragraph."""
    para = anchor.insert_paragraph_before("")
    _add_marked_run(para, text, "delete")


def _append_deleted_paragraph(doc, text: str) -> None:
    """Append a struck-through paragraph at the end of the document.

    Used when the removed text ran past the last surviving paragraph, where
    there is nothing left to anchor an insertion against.
    """
    _add_marked_run(doc.add_paragraph(""), text, "delete")


def _record_deletions(doc, new_paras, anchor_index: int, texts) -> int:
    """Mark removed paragraphs, anchoring them or appending past the end."""
    anchor = new_paras[anchor_index] if anchor_index < len(new_paras) else None
    removed = 0
    for text in texts:
        if not text.strip():
            continue
        if anchor is not None:
            _insert_deleted_paragraph(anchor, text)
        else:
            _append_deleted_paragraph(doc, text)
        removed += 1
    return removed


def mark_docx_changes(
    old_docx: Path,
    new_docx: Path,
    output_path: Path,
) -> Tuple[Path, dict]:
    """Write a copy of ``new_docx`` with changes against ``old_docx`` marked.

    Args:
        old_docx: The baseline export (for example, the submitted version)
        new_docx: The revised export
        output_path: Where to write the marked-up DOCX

    Returns:
        The output path and a dict of change counts
    """
    old_doc = Document(str(old_docx))
    new_doc = Document(str(new_docx))

    old_paras = _iter_block_paragraphs(old_doc)
    new_paras = _iter_block_paragraphs(new_doc)
    old_texts = [p.text for p in old_paras]
    new_texts = [p.text for p in new_paras]

    matcher = difflib.SequenceMatcher(None, old_texts, new_texts, autojunk=False)
    stats = {"unchanged": 0, "changed": 0, "inserted": 0, "deleted": 0}

    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == "equal":
            stats["unchanged"] += j2 - j1
            continue

        if tag == "insert":
            for para in new_paras[j1:j2]:
                if para.text.strip():
                    _mark_paragraph(para, "", para.text)
                    stats["inserted"] += 1
            continue

        if tag == "delete":
            stats["deleted"] += _record_deletions(new_doc, new_paras, j1, old_texts[i1:i2])
            continue

        # replace: pair paragraphs positionally, then handle any leftovers
        paired = min(i2 - i1, j2 - j1)
        for offset in range(paired):
            old_text, para = old_texts[i1 + offset], new_paras[j1 + offset]
            if old_text == para.text:
                continue
            if _similarity(old_text, para.text) >= SIMILARITY_THRESHOLD:
                _mark_paragraph(para, old_text, para.text)
                stats["changed"] += 1
            else:
                # Unrelated paragraphs: show the old one removed and the new one added.
                if old_text.strip():
                    _insert_deleted_paragraph(para, old_text)
                    stats["deleted"] += 1
                _mark_paragraph(para, "", para.text)
                stats["inserted"] += 1
        for para in new_paras[j1 + paired : j2]:
            if para.text.strip():
                _mark_paragraph(para, "", para.text)
                stats["inserted"] += 1
        stats["deleted"] += _record_deletions(new_doc, new_paras, j1 + paired, old_texts[i1 + paired : i2])

    output_path.parent.mkdir(parents=True, exist_ok=True)
    new_doc.save(str(output_path))
    logger.info(
        f"Marked changes: {stats['changed']} paragraph(s) modified, "
        f"{stats['inserted']} added, {stats['deleted']} removed"
    )
    return output_path, stats


def build_tracked_docx(
    manuscript_path: Path,
    git_tag: str,
    exporter_factory,
    output_path: Optional[Path] = None,
) -> Tuple[Path, dict]:
    """Export the manuscript twice and mark what changed since ``git_tag``.

    Args:
        manuscript_path: Manuscript directory in the working tree
        git_tag: Tag identifying the baseline (for example the submitted version)
        exporter_factory: Callable taking a manuscript path and returning a
            configured exporter with an ``export()`` method
        output_path: Where to write the marked file; defaults to a sibling of the
            normal export named ``__changes_vs_<tag>``

    Returns:
        The output path and a dict of change counts
    """
    import shutil
    import subprocess
    import tempfile

    manuscript_path = Path(manuscript_path).resolve()
    repo_root = manuscript_path.parent

    with tempfile.TemporaryDirectory() as tmp:
        tag_root = Path(tmp) / "baseline"
        tag_root.mkdir()

        archive = subprocess.Popen(
            ["git", "archive", "--format=tar", git_tag],
            stdout=subprocess.PIPE,
            cwd=repo_root,
        )
        subprocess.check_call(["tar", "-x", "-C", str(tag_root)], stdin=archive.stdout)
        archive.wait()
        if archive.returncode != 0:
            raise RuntimeError(f"Could not extract git tag '{git_tag}'")

        tag_manuscript = tag_root / manuscript_path.name
        if not tag_manuscript.exists():
            raise RuntimeError(f"Manuscript directory '{manuscript_path.name}' is not present at tag '{git_tag}'")

        baseline_docx = Path(tmp) / "baseline.docx"
        shutil.copy(exporter_factory(tag_manuscript).export(), baseline_docx)

        revised_docx = exporter_factory(manuscript_path).export()
        if output_path is None:
            safe_tag = re.sub(r"[^A-Za-z0-9._-]+", "-", git_tag)
            output_path = revised_docx.with_name(f"{revised_docx.stem}__changes_vs_{safe_tag}.docx")

        return mark_docx_changes(baseline_docx, revised_docx, output_path)
