"""Tests for DOCX revision marking."""

from pathlib import Path

import pytest
from docx import Document

from rxiv_maker.exporters.docx_diff import mark_docx_changes


def _write(tmp_path: Path, name: str, paragraphs) -> Path:
    doc = Document()
    for text in paragraphs:
        doc.add_paragraph(text)
    path = tmp_path / name
    doc.save(str(path))
    return path


def _mark(tmp_path: Path, before, after):
    old = _write(tmp_path, "old.docx", before)
    new = _write(tmp_path, "new.docx", after)
    out, stats = mark_docx_changes(old, new, tmp_path / "marked.docx")
    return Document(str(out)), stats


def _struck(doc):
    return [r.text for p in doc.paragraphs for r in p.runs if r.font.strike]


def _highlighted(doc):
    return [r.text for p in doc.paragraphs for r in p.runs if r.font.highlight_color and not r.font.strike]


def _plain(doc):
    return [r.text for p in doc.paragraphs for r in p.runs if not r.font.strike and not r.font.highlight_color]


class TestWordLevelMarking:
    """Small edits inside a paragraph are marked word by word."""

    def test_single_word_edit_is_marked(self, tmp_path):
        doc, stats = _mark(tmp_path, ["The quick brown fox."], ["The quick red fox."])
        assert "brown " in _struck(doc)
        assert "red " in _highlighted(doc)
        assert stats["changed"] == 1

    def test_untouched_words_stay_plain(self, tmp_path):
        doc, _ = _mark(tmp_path, ["The quick brown fox."], ["The quick red fox."])
        plain = "".join(_plain(doc))
        assert "The quick " in plain
        assert "fox." in plain

    def test_identical_document_marks_nothing(self, tmp_path):
        doc, stats = _mark(tmp_path, ["Same line.", "Second line."], ["Same line.", "Second line."])
        assert _struck(doc) == []
        assert _highlighted(doc) == []
        assert stats["changed"] == 0

    def test_citation_style_edit_is_marked(self, tmp_path):
        """The real case: nested parens collapsed to a proper author-date group."""
        doc, _ = _mark(
            tmp_path,
            ["Shown by Brito et al. (2025); Grobe et al. (2026)) here."],
            ["Shown by Brito et al., 2025; Grobe et al., 2026) here."],
        )
        assert any("(2025)" in t for t in _struck(doc))
        assert any("2025;" in t for t in _highlighted(doc))


class TestParagraphLevelMarking:
    """Whole paragraphs added or removed are shown as such."""

    def test_added_paragraph_is_highlighted(self, tmp_path):
        doc, stats = _mark(tmp_path, ["Opening."], ["Opening.", "A wholly new sentence appears."])
        assert "A wholly new sentence appears." in _highlighted(doc)
        assert stats["inserted"] == 1

    def test_deleted_paragraph_is_struck_through(self, tmp_path):
        doc, stats = _mark(tmp_path, ["Opening.", "This goes away."], ["Opening."])
        assert stats["deleted"] == 1 or "This goes away." in _struck(doc)

    def test_unlike_paragraphs_split_instead_of_interleaving(self, tmp_path):
        """Unrelated paragraphs must not be word-diffed into a hybrid."""
        doc, _ = _mark(
            tmp_path,
            ["Alpha beta gamma delta epsilon.", "Tail."],
            ["Completely different wording entirely.", "Tail."],
        )
        texts = [p.text for p in doc.paragraphs]
        assert "Alpha beta gamma delta epsilon." in texts
        assert "Completely different wording entirely." in texts

    def test_close_paragraphs_are_word_diffed(self, tmp_path):
        """Above the similarity threshold, keep the fine-grained diff."""
        doc, stats = _mark(
            tmp_path,
            ["The framework supports Python and R scripts."],
            ["The framework supports Python and Julia scripts."],
        )
        assert stats["changed"] == 1
        assert any("Julia" in t for t in _highlighted(doc))


class TestOutput:
    """The marked file is a real, openable document."""

    def test_output_file_is_written(self, tmp_path):
        old = _write(tmp_path, "old.docx", ["One."])
        new = _write(tmp_path, "new.docx", ["Two."])
        out, _ = mark_docx_changes(old, new, tmp_path / "sub" / "marked.docx")
        assert out.exists()
        Document(str(out))

    def test_stats_include_every_category(self, tmp_path):
        _, stats = _mark(tmp_path, ["a"], ["a"])
        assert set(stats) == {"unchanged", "changed", "inserted", "deleted"}


@pytest.mark.parametrize("threshold_case", ["identical", "small edit"])
def test_threshold_cases_do_not_crash(tmp_path, threshold_case):
    before = ["Shared opening line."]
    after = ["Shared opening line."] if threshold_case == "identical" else ["Shared opening lines."]
    mark_docx_changes(
        _write(tmp_path, "o.docx", before),
        _write(tmp_path, "n.docx", after),
        tmp_path / "m.docx",
    )
